from docx import  Document
from docx.shared import Inches
import glob
from os import listdir
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def center_insert_img(doc,img):
    """插入图片"""
    for paragraph in doc.paragraphs:
        # 根据文档中占位符定位图片插入的位置
        if '<<img1>>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<<img1>>', '')
            run = paragraph.add_run('')
            run.add_break()

            run.add_picture(img, width=Inches(6.2))


def save_img_to_doc(img):
    tpl_doc = r'D:\py\归档模版.docx'
    res_doc = r'D:\py\空港.docx'

    document = Document(tpl_doc)

    center_insert_img(document,img)

    document.save(res_doc)


def main():
    path = r'C:\Users\durpa\Desktop\身份证\*.jpg'
    # img = 'C:\\Users\\durpa\\Desktop\\身份证\\75+'
    for img in glob.glob(path):
        print(img)



    save_img_to_doc(img)


if __name__ == '__main__':
    main()